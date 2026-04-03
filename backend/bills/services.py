"""
Bill processing services for BillClear AI.

This module contains all AI-powered logic for parsing, analyzing, and chatting
about medical bills using the Anthropic Claude API.
"""

import json
import base64
import logging

import anthropic
from django.conf import settings

logger = logging.getLogger(__name__)

MODEL = "claude-sonnet-4-6"


def _get_client() -> anthropic.Anthropic:
    """Return a configured Anthropic client."""
    return anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)


def parse_bill(bill_instance) -> dict:
    """
    Parse a medical bill file using Claude's vision capability.

    Sends the uploaded bill image/PDF to Claude with a structured extraction
    prompt. Returns a dict of extracted data and updates the Bill model fields.

    Args:
        bill_instance: A bills.models.Bill instance with an uploaded original_file.

    Returns:
        A dict containing the parsed bill data as extracted by Claude.
    """
    client = _get_client()

    # Read the uploaded file and encode it for the API
    bill_instance.original_file.open("rb")
    file_bytes = bill_instance.original_file.read()
    bill_instance.original_file.close()

    file_name = bill_instance.original_file.name.lower()
    if file_name.endswith(".pdf"):
        media_type = "application/pdf"
    elif file_name.endswith(".png"):
        media_type = "image/png"
    elif file_name.endswith((".jpg", ".jpeg")):
        media_type = "image/jpeg"
    else:
        media_type = "image/jpeg"

    encoded_file = base64.standard_b64encode(file_bytes).decode("utf-8")

    system_prompt = """You are a medical billing expert. Your task is to extract structured data
from a medical bill or Explanation of Benefits (EOB) document.

Return ONLY valid JSON with this exact structure:
{
  "provider_name": "string",
  "facility_type": "hospital|clinic|lab|pharmacy|specialist|other",
  "date_of_service": "YYYY-MM-DD or null",
  "total_charged": number or null,
  "total_allowed": number or null,
  "patient_responsibility": number or null,
  "line_items": [
    {
      "cpt_code": "5-digit string or null",
      "hcpcs_code": "string or null",
      "icd10_codes": ["list", "of", "codes"],
      "description_raw": "exact text from bill",
      "quantity": integer,
      "charged_amount": number,
      "allowed_amount": number or null
    }
  ]
}

Rules:
- CPT codes must be exactly 5 digits. Set to null if not clearly identifiable.
- Extract ALL line items visible on the bill.
- If a field is not present, use null.
- Do not include any text outside the JSON object."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": encoded_file,
                        },
                    },
                    {
                        "type": "text",
                        "text": "Please extract all billing information from this medical bill.",
                    },
                ],
            }
        ],
    )

    raw_text = response.content[0].text.strip()

    # Strip markdown code fences if present
    if raw_text.startswith("```"):
        raw_text = raw_text.split("\n", 1)[1]
        raw_text = raw_text.rsplit("```", 1)[0]

    parsed_data = json.loads(raw_text)

    # Update bill instance fields
    bill_instance.provider_name = parsed_data.get("provider_name") or ""
    bill_instance.facility_type = parsed_data.get("facility_type") or "other"
    bill_instance.total_charged = parsed_data.get("total_charged")
    bill_instance.total_allowed = parsed_data.get("total_allowed")
    bill_instance.patient_responsibility = parsed_data.get("patient_responsibility")

    date_str = parsed_data.get("date_of_service")
    if date_str:
        from datetime import date
        try:
            bill_instance.date_of_service = date.fromisoformat(date_str)
        except ValueError:
            bill_instance.date_of_service = None

    bill_instance.save()

    return parsed_data


def analyze_line_items(bill_instance) -> None:
    """
    Analyze line items on a bill for errors, overcharges, and risk levels.

    Uses Claude to translate CPT/HCPCS codes into plain English, detect
    potential billing errors (duplicates, upcoding, unbundling, balance billing),
    and assign a risk level (green/yellow/red) to each line item.

    Args:
        bill_instance: A bills.models.Bill instance with related LineItem objects.
    """
    client = _get_client()

    line_items = list(bill_instance.line_items.all())
    if not line_items:
        logger.warning("No line items found for bill %s", bill_instance.pk)
        return

    items_data = [
        {
            "id": item.pk,
            "cpt_code": item.cpt_code,
            "hcpcs_code": item.hcpcs_code,
            "description_raw": item.description_raw,
            "quantity": item.quantity,
            "charged_amount": float(item.charged_amount),
            "allowed_amount": float(item.allowed_amount) if item.allowed_amount else None,
        }
        for item in line_items
    ]

    system_prompt = """You are a medical billing auditor with expertise in CPT codes, HCPCS codes,
and common medical billing errors.

Analyze the provided line items and return ONLY valid JSON with this structure:
{
  "line_items": [
    {
      "id": integer,
      "description_plain": "plain English explanation of what this charge is for",
      "risk_level": "green|yellow|red",
      "error_type": "duplicate|unbundled|upcoded|balance_billing|null",
      "flag_explanation": "string — if risk_level is yellow or red, explain specifically why this charge was flagged and what the billing issue means for this particular charge. Be concrete and patient-friendly. Empty string if risk_level is green."
    }
  ]
}

Risk level guide:
- green: Charge appears reasonable and correct
- yellow: Charge warrants review (unusual quantity, slightly high, unclear code)
- red: Likely error or overcharge (duplicate, known upcoding pattern, balance billing)

Error types:
- duplicate: Same service billed more than once
- unbundled: Services that should be billed together are split to inflate cost
- upcoded: Service billed at a higher complexity/cost than documented
- balance_billing: Patient charged more than the allowed amount
- null: No apparent error

Return null for error_type and empty string for flag_explanation when risk_level is green."""

    response = client.messages.create(
        model=MODEL,
        max_tokens=4096,
        system=system_prompt,
        messages=[
            {
                "role": "user",
                "content": json.dumps({"line_items": items_data}),
            }
        ],
    )

    raw_text = response.content[0].text.strip()
    if raw_text.startswith("```"):
        raw_text = raw_text.split("\n", 1)[1]
        raw_text = raw_text.rsplit("```", 1)[0]

    analysis = json.loads(raw_text)

    # Build a lookup map for quick access
    analysis_map = {item["id"]: item for item in analysis.get("line_items", [])}

    for item in line_items:
        if item.pk in analysis_map:
            result = analysis_map[item.pk]
            item.description_plain = result.get("description_plain", "")
            item.risk_level = result.get("risk_level", "green")
            item.error_type = result.get("error_type") or None
            item.flag_explanation = result.get("flag_explanation") or ""
            item.save()

    # Update bill status from new to reviewed
    if bill_instance.status == "new":
        bill_instance.status = "reviewed"
        bill_instance.save(update_fields=["status"])


def generate_dispute_letter(dispute_instance) -> None:
    """
    Generate a formal dispute letter as a Word document (.docx) and plain-text preview.

    Builds a properly formatted Word document using python-docx with real headings,
    bold text, and a table for disputed charges. Saves the .docx to letter_pdf and
    a plain-text version to letter_content for on-screen preview.

    Auto-fills all available user and bill data. Only leaves brackets for:
    - [ACCOUNT/PATIENT ID NUMBER]
    - [EOB DATE]

    Args:
        dispute_instance: A disputes.models.Dispute instance with related line items.
    """
    from io import BytesIO
    from datetime import date as _date
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from django.core.files.base import ContentFile

    bill = dispute_instance.bill
    user = bill.user
    line_items = list(dispute_instance.line_items.all())

    # ── Resolve sender fields ──────────────────────────────────────────────────
    patient_name = f"{user.first_name} {user.last_name}".strip() or "[YOUR FULL NAME]"
    patient_email = user.email
    patient_phone = user.phone_number or "[YOUR PHONE NUMBER]"
    street = user.street_address or "[YOUR STREET ADDRESS]"
    city_state_zip_parts = []
    if user.city:
        city_state_zip_parts.append(user.city)
    if user.state:
        city_state_zip_parts.append(user.state)
    if user.zip_code:
        city_state_zip_parts.append(user.zip_code)
    city_state_zip = ", ".join(city_state_zip_parts) if city_state_zip_parts else "[CITY, STATE ZIP]"

    today_str = _date.today().strftime("%B %d, %Y")
    provider = bill.provider_name or "Healthcare Provider"
    date_of_service = (
        bill.date_of_service.strftime("%B %d, %Y") if bill.date_of_service else "Date on File"
    )
    insurance_provider = user.insurance_provider or "Insurance Provider on File"
    plan_type = user.plan_type or ""

    # ── Build the Word document ────────────────────────────────────────────────
    doc = Document()

    # Default font / margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def _add_para(text="", bold=False, italic=False, size=11, space_after=6, alignment=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        if alignment is not None:
            p.alignment = alignment
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.size = Pt(size)
        return p

    def _add_run(para, text, bold=False, italic=False, size=11):
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return run

    # ── Instructions (italic, shaded) ─────────────────────────────────────────
    p = _add_para(
        "Fields marked with brackets [ ] need to be filled in before sending this letter.",
        italic=True,
        size=10,
        space_after=12,
    )
    # Light yellow shading on the instruction paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF9C4")
    pPr.append(shd)

    # ── Sender block ──────────────────────────────────────────────────────────
    _add_para(patient_name, bold=True, space_after=2)
    _add_para(street, space_after=2)
    _add_para(city_state_zip, space_after=2)
    _add_para(patient_phone, space_after=2)
    _add_para(patient_email, space_after=12)

    # ── Date ──────────────────────────────────────────────────────────────────
    _add_para(today_str, space_after=12)

    # ── Recipient block ───────────────────────────────────────────────────────
    _add_para("Billing Department", bold=True, space_after=2)
    _add_para(provider, space_after=2)
    _add_para("Address on File", space_after=12)

    # ── RE line ───────────────────────────────────────────────────────────────
    re_para = doc.add_paragraph()
    re_para.paragraph_format.space_after = Pt(12)
    _add_run(re_para, "RE: ", bold=True, size=11)
    _add_run(
        re_para,
        f"Formal Dispute of Medical Bill – Date of Service: {date_of_service}",
        bold=True,
        size=11,
    )

    # ── Salutation ────────────────────────────────────────────────────────────
    _add_para("Dear Billing Department,", space_after=10)

    # ── Opening paragraph ─────────────────────────────────────────────────────
    _add_para(
        f"I am writing to formally dispute the following charges appearing on my medical bill "
        f"from {provider} for services rendered on {date_of_service}. After carefully reviewing "
        f"my billing statement, I have identified potential errors or overcharges that I believe "
        f"require your immediate attention and correction.",
        space_after=10,
    )

    # ── Insurance / account context ───────────────────────────────────────────
    ins_line = f"My insurance provider is {insurance_provider}"
    if plan_type:
        ins_line += f" ({plan_type})"
    ins_line += (
        ". My account/patient ID number is [ACCOUNT/PATIENT ID NUMBER] and the relevant "
        "EOB date is [EOB DATE]."
    )
    _add_para(ins_line, space_after=10)

    # ── Disputed charges heading ──────────────────────────────────────────────
    _add_para("Disputed Charges:", bold=True, space_after=6)

    # ── Table ─────────────────────────────────────────────────────────────────
    col_widths = [Inches(1), Inches(2.4), Inches(1.0), Inches(1.0), Inches(1.4)]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = ["Code", "Service Description", "Charged", "Allowed", "Issue Identified"]
    for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
        cell.width = col_widths[i]
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        # Header background
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "DBEAFE")
        tc_pr.append(shd)

    for item in line_items:
        row_cells = table.add_row().cells
        code = item.cpt_code or item.hcpcs_code or "N/A"
        desc = item.description_plain or item.description_raw or "N/A"
        charged = f"${item.charged_amount:,.2f}"
        allowed = f"${item.allowed_amount:,.2f}" if item.allowed_amount is not None else "N/A"
        issue = item.error_type.replace("_", " ").title() if item.error_type else "Under Review"
        for cell, text in zip(row_cells, [code, desc, charged, allowed, issue]):
            cell.paragraphs[0].add_run(text).font.size = Pt(10)

    doc.add_paragraph()  # spacer after table

    # ── Flag explanations (if any) ────────────────────────────────────────────
    flagged = [i for i in line_items if i.flag_explanation]
    if flagged:
        _add_para("Additional details on identified issues:", bold=True, space_after=4)
        for item in flagged:
            code = item.cpt_code or item.hcpcs_code or "N/A"
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_after = Pt(4)
            _add_run(bullet, f"{code}: ", bold=True, size=10)
            _add_run(bullet, item.flag_explanation, size=10)
        doc.add_paragraph()

    # ── Billing summary ───────────────────────────────────────────────────────
    _add_para("Billing Summary from Statement:", bold=True, space_after=4)
    summary_lines = []
    if bill.total_charged is not None:
        summary_lines.append(f"Total Charged: ${bill.total_charged:,.2f}")
    if bill.total_allowed is not None:
        summary_lines.append(f"Total Allowed: ${bill.total_allowed:,.2f}")
    if bill.patient_responsibility is not None:
        summary_lines.append(f"Patient Responsibility: ${bill.patient_responsibility:,.2f}")
    for line in summary_lines:
        b = doc.add_paragraph(style="List Bullet")
        b.paragraph_format.space_after = Pt(3)
        b.add_run(line).font.size = Pt(10)
    if summary_lines:
        doc.add_paragraph()

    # ── Requests ──────────────────────────────────────────────────────────────
    _add_para(
        "In light of the above, I formally request the following:", space_after=4
    )
    requests = [
        "A written explanation for each disputed charge listed in the table above.",
        "A corrected, itemized statement reflecting any applicable adjustments or corrections.",
        "A refund of any amount I have paid in excess of what is lawfully owed.",
        "Confirmation that my account will not be sent to collections while this dispute is under review.",
    ]
    for req in requests:
        b = doc.add_paragraph(style="List Bullet")
        b.paragraph_format.space_after = Pt(4)
        b.add_run(req).font.size = Pt(11)
    doc.add_paragraph()

    # ── Closing paragraph ─────────────────────────────────────────────────────
    _add_para(
        f"Please respond in writing within 30 days of receiving this letter. "
        f"You may reach me by email at {patient_email} or by phone at {patient_phone}.",
        space_after=12,
    )
    _add_para("Sincerely,", space_after=36)  # space for handwritten signature

    # ── Signature block ───────────────────────────────────────────────────────
    _add_para(patient_name, bold=True, space_after=2)
    _add_para(patient_email, space_after=2)
    if user.phone_number:
        _add_para(patient_phone, space_after=0)

    # ── Save docx to letter_pdf FileField ─────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"dispute-letter-{dispute_instance.pk}.docx"
    dispute_instance.letter_pdf.save(filename, ContentFile(buffer.read()), save=False)

    # ── Build plain-text preview for letter_content ───────────────────────────
    lines = [
        "NOTE: Download the Word document below for the fully formatted letter.",
        "",
        patient_name,
        street,
        city_state_zip,
        patient_phone,
        patient_email,
        "",
        today_str,
        "",
        "Billing Department",
        provider,
        "",
        f"RE: Formal Dispute of Medical Bill – Date of Service: {date_of_service}",
        "",
        "Dear Billing Department,",
        "",
        f"I am writing to formally dispute the following charges from {provider} for services "
        f"rendered on {date_of_service}. My insurance provider is {ins_line.split('. My')[0].replace('My insurance provider is ', '')}.",
        "",
        "Disputed Charges:",
    ]
    for item in line_items:
        code = item.cpt_code or item.hcpcs_code or "N/A"
        desc = item.description_plain or item.description_raw or "N/A"
        charged = f"${item.charged_amount:,.2f}"
        issue = item.error_type.replace("_", " ").title() if item.error_type else "Under Review"
        lines.append(f"  • {code} – {desc} | Charged: {charged} | Issue: {issue}")
        if item.flag_explanation:
            lines.append(f"    {item.flag_explanation}")
    lines += [
        "",
        "Account/Patient ID: [ACCOUNT/PATIENT ID NUMBER]",
        "EOB Date: [EOB DATE]",
        "",
        "I request a written explanation, corrected itemized statement, and refund of overcharged amounts. "
        "Please respond within 30 days.",
        "",
        "Sincerely,",
        "",
        patient_name,
        patient_email,
    ]
    letter_content = "\n".join(lines)

    dispute_instance.letter_content = letter_content
    dispute_instance.save(update_fields=["letter_content", "letter_pdf"])


def chat_with_bill(bill_instance, user_message: str) -> str:
    """
    Respond to a user's question about a specific medical bill.

    Provides context about the bill's line items and any identified issues,
    then answers the user's question in plain, helpful language.

    Args:
        bill_instance: A bills.models.Bill instance to provide context.
        user_message: The user's question or message as a string.

    Returns:
        The assistant's response as a string.
    """
    client = _get_client()

    line_items = list(bill_instance.line_items.all())
    user = bill_instance.user

    def _fmt_item(item):
        code = item.cpt_code or item.hcpcs_code or "N/A"
        line = (
            f"  Code: {code} | {item.description_plain or item.description_raw} | "
            f"Charged: ${item.charged_amount}"
        )
        if item.allowed_amount is not None:
            line += f" | Allowed: ${item.allowed_amount}"
        if item.regional_average is not None:
            line += f" | Regional avg: ${item.regional_average}"
        line += f" | Risk: {item.risk_level}"
        if item.error_type:
            line += f" | Issue: {item.error_type}"
        if item.flag_explanation:
            line += f" | Explanation: {item.flag_explanation}"
        return line

    items_context = "\n".join(_fmt_item(item) for item in line_items)

    zip_info = f"Patient ZIP code: {user.zip_code}" if getattr(user, "zip_code", None) else "Patient ZIP code: not provided"
    insurance_info = f"Insurance: {user.insurance_provider}" if getattr(user, "insurance_provider", None) else "Insurance: not provided"

    # Fetch conversation history
    from bills.models import ChatMessage

    history = ChatMessage.objects.filter(bill=bill_instance).order_by("created_at")
    conversation = [
        {"role": msg.role, "content": msg.content}
        for msg in history
    ]

    system_prompt = f"""You are a helpful medical billing assistant for BillClear AI.
You help patients understand their medical bills, identify potential errors, and take action.

Bill context:
- Provider: {bill_instance.provider_name or "Unknown"}
- Date of Service: {bill_instance.date_of_service or "Unknown"}
- Total Charged: ${bill_instance.total_charged or "Unknown"}
- Total Allowed: ${bill_instance.total_allowed or "Unknown"}
- Patient Responsibility: ${bill_instance.patient_responsibility or "Unknown"}
- Status: {bill_instance.status}
- {zip_info}
- {insurance_info}

Line Items (code | description | amounts | risk):
{items_context if items_context else "No line items have been extracted yet."}

Use the line item data above — including codes, amounts, regional averages, and flagged issues — to give specific, contextual answers about this bill.
Answer in plain language a non-expert can understand. When relevant, cite the specific code or charge you're discussing.
Always recommend the patient verify information with their provider or insurer."""

    # Append the new user message
    conversation.append({"role": "user", "content": user_message})

    response = client.messages.create(
        model=MODEL,
        max_tokens=1024,
        system=system_prompt,
        messages=conversation,
    )

    return response.content[0].text.strip()
